#!/usr/bin/env python3
"""Fail-closed structural and package parity audit for AI Security Manuals.

Protected-token checks intentionally cover identifiers that should ordinarily
remain unchanged in translation: standards, protocols, security frameworks,
product/platform names, technical acronyms, and URLs. Generic all-caps English
words, translated concepts such as AI/IA, and repeated section numbers are not
used as token-parity evidence.
"""
from __future__ import annotations

import hashlib
import json
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path
from urllib.parse import unquote

from docx import Document
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
MANUALS = ROOT / "manuals"
QA = ROOT / "qa" / "full-multilingual-parity"
QA.mkdir(parents=True, exist_ok=True)
LANGS = ("en", "es-419", "pt-BR")

# Curated immutable or ordinarily untranslated technical identifiers.
PROTECTED_NAMES = (
    "ISO/IEC", "NIST", "OWASP", "MITRE", "ATLAS", "AWS", "Azure",
    "Google Cloud", "Oracle Cloud", "IBM Cloud", "CVE", "API", "CLI",
    "IAM", "MFA", "TLS", "HTTP", "HTTPS", "TCP", "UDP", "JSON",
    "YAML", "XML", "PDF", "DOCX", "RAG", "MCP", "LLM", "ML",
    "SQL", "XSS", "PII", "DPIA", "RACI", "AICM", "CSA", "CISA",
    "OECD", "UNESCO", "CI/CD", "Kubernetes", "Terraform", "Ansible",
)
_NAME_PATTERN = "|".join(sorted((re.escape(x) for x in PROTECTED_NAMES), key=len, reverse=True))
PROTECTED = re.compile(rf"https?://[^\s)>]+|(?<![\w/])(?:{_NAME_PATTERN})(?![\w/])", re.IGNORECASE)


def canonical_token(value: str) -> str:
    value = value.rstrip(".,;:]")
    if value.lower().startswith(("http://", "https://")):
        return value
    return value.casefold()


def token_set(text: str) -> set[str]:
    return {canonical_token(x) for x in PROTECTED.findall(text)}


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def docx_info(path: Path) -> dict:
    out = {"valid": False, "paragraphs": 0, "words": 0, "headings": [], "tables": 0, "images": 0, "tokens": []}
    try:
        with zipfile.ZipFile(path) as package:
            bad = package.testzip()
            if bad:
                out["error"] = f"corrupt ZIP member: {bad}"
                return out
            out["images"] = sum(1 for name in package.namelist() if name.startswith("word/media/"))
        document = Document(path)
        texts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            texts.append(text)
            if paragraph.style and paragraph.style.name.lower().startswith("heading"):
                out["headings"].append({"style": paragraph.style.name, "text": text})
        for table in document.tables:
            for row in table.rows:
                texts.extend(cell.text.strip() for cell in row.cells if cell.text.strip())
        full = "\n".join(texts)
        out.update(
            valid=True,
            paragraphs=len(texts),
            words=len(re.findall(r"\b\w+\b", full)),
            tables=len(document.tables),
            tokens=sorted(token_set(full)),
        )
    except Exception as exc:
        out["error"] = repr(exc)
    return out


def pdf_info(path: Path) -> dict:
    out = {"valid": False, "pages": 0, "words": 0, "blank_pages": [], "tokens": []}
    try:
        reader = PdfReader(str(path))
        texts: list[str] = []
        for number, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            if len(text.strip()) < 20:
                out["blank_pages"].append(number)
            texts.append(text)
        full = "\n".join(texts)
        out.update(
            valid=True,
            pages=len(reader.pages),
            words=len(re.findall(r"\b\w+\b", full)),
            tokens=sorted(token_set(full)),
        )
    except Exception as exc:
        out["error"] = repr(exc)
    return out


def local_links() -> list[dict]:
    findings: list[dict] = []
    link_pattern = re.compile(r"!?\[[^\]]*\]\(([^)]+)\)")
    for markdown in ROOT.rglob("*.md"):
        text = markdown.read_text(encoding="utf-8", errors="replace")
        if "\ufffd" in text:
            findings.append({"severity": "HIGH", "file": str(markdown.relative_to(ROOT)), "finding": "Unicode replacement character"})
        for raw in link_pattern.findall(text):
            target = raw.split("#", 1)[0].strip().strip("<>")
            if not target or re.match(r"^(?:https?://|mailto:)", target):
                continue
            destination = (markdown.parent / unquote(target)).resolve()
            if ROOT.resolve() not in destination.parents and destination != ROOT.resolve():
                findings.append({"severity": "HIGH", "file": str(markdown.relative_to(ROOT)), "finding": f"link escapes repository: {raw}"})
            elif not destination.exists():
                findings.append({"severity": "BLOCKER", "file": str(markdown.relative_to(ROOT)), "finding": f"broken local link: {raw}"})
    return findings


def family_dirs() -> list[Path]:
    found: list[Path] = []
    for english in MANUALS.rglob("en"):
        if english.is_dir() and all((english.parent / language).is_dir() for language in LANGS):
            found.append(english.parent)
    return sorted(set(found))


def one_file(folder: Path, suffix: str) -> Path | None:
    files = sorted(folder.glob(f"*{suffix}"))
    return files[0] if len(files) == 1 else None


def main() -> int:
    inventory: list[dict] = []
    findings = local_links()
    families: list[dict] = []

    for family in family_dirs():
        record = {"family": str(family.relative_to(ROOT)), "languages": {}, "status": "PASS"}
        for language in LANGS:
            folder = family / language
            language_record: dict = {}
            for extension in (".docx", ".pdf"):
                path = one_file(folder, extension)
                if not path:
                    findings.append({"severity": "BLOCKER", "file": str(folder.relative_to(ROOT)), "finding": f"expected exactly one {extension} file"})
                    continue
                info = docx_info(path) if extension == ".docx" else pdf_info(path)
                info.update(path=str(path.relative_to(ROOT)), sha256=sha256(path), bytes=path.stat().st_size)
                language_record[extension[1:]] = info
                inventory.append({"family": record["family"], "language": language, "format": extension[1:], **info})
                if not info.get("valid"):
                    findings.append({"severity": "BLOCKER", "file": str(path.relative_to(ROOT)), "finding": info.get("error", "package invalid")})
            record["languages"][language] = language_record

        english = record["languages"].get("en", {}).get("docx", {})
        english_headings = [heading["style"] for heading in english.get("headings", [])]
        english_tokens = set(english.get("tokens", []))

        for language in ("es-419", "pt-BR"):
            localized = record["languages"].get(language, {}).get("docx", {})
            localized_headings = [heading["style"] for heading in localized.get("headings", [])]
            if english_headings != localized_headings:
                findings.append({
                    "severity": "HIGH",
                    "file": record["family"],
                    "finding": f"{language} heading-style sequence differs from English",
                    "english_count": len(english_headings),
                    "localized_count": len(localized_headings),
                })

            english_words = english.get("words", 0)
            localized_words = localized.get("words", 0)
            ratio = localized_words / english_words if english_words else 0
            if not 0.70 <= ratio <= 1.45:
                findings.append({"severity": "HIGH", "file": record["family"], "finding": f"{language} DOCX word-count ratio outside 0.70-1.45", "ratio": round(ratio, 3)})

            localized_tokens = set(localized.get("tokens", []))
            missing = sorted(english_tokens - localized_tokens)
            if missing:
                findings.append({
                    "severity": "HIGH",
                    "file": record["family"],
                    "finding": f"{language} missing curated protected identifiers found in English",
                    "examples": missing[:25],
                    "count": len(missing),
                })
        families.append(record)

    severity = Counter(finding["severity"] for finding in findings)
    overall = "FAIL" if severity["BLOCKER"] or severity["HIGH"] else "BLOCKED PENDING HUMAN REVIEW"
    summary = {
        "status": overall,
        "family_count": len(families),
        "edition_count": len(inventory),
        "findings_by_severity": dict(severity),
        "human_review_required": True,
        "limitation": "Automated structural/package/identifier checks cannot certify sentence-level semantic equivalence or native-language editorial approval.",
    }

    (QA / "FULL_MULTILINGUAL_CONTENT_PARITY_INVENTORY.json").write_text(json.dumps(inventory, indent=2, ensure_ascii=False), encoding="utf-8")
    (QA / "FULL_MULTILINGUAL_CONTENT_PARITY_FINDINGS.json").write_text(json.dumps(findings, indent=2, ensure_ascii=False), encoding="utf-8")
    lines = [
        "# Full Multilingual Content Parity Audit", "", f"**Status:** {overall}", "",
        f"Manual families: {len(families)}  ", f"Publication files audited: {len(inventory)}  ",
        f"Findings: {dict(severity)}", "", "## Scope and conclusion", "",
        "The audit validates repository links, DOCX ZIP/XML readability, PDF text extraction, package inventories, heading-style parity, word-count ratios, and curated protected-identifier presence across English, Latin American Spanish, and Brazilian Portuguese editions.",
        "", "Generic translated terms, arbitrary all-caps words, repetition counts, and section numbers are excluded from protected-identifier findings.",
        "", "It does **not** claim native-language or sentence-level semantic approval. Every family remains blocked pending documented human editorial review unless separate exact-SHA evidence exists.",
        "", "## Findings",
    ]
    if findings:
        lines.extend(f"- **{finding['severity']}** — `{finding['file']}` — {finding['finding']}" for finding in findings)
    else:
        lines.append("- No automated blocker or high-severity defects detected.")
    lines += ["", "## Release rule", "", "Do not mark an edition fully approved until automated defects are resolved and English technical, Spanish native-language, Portuguese native-language, accessibility, and page-by-page visual reviews are documented at the exact candidate SHA."]
    (QA / "FULL_MULTILINGUAL_CONTENT_PARITY_AUDIT.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    checksums = [f"{row['sha256']}  {row['path']}" for row in inventory]
    (QA / "FULL_MULTILINGUAL_SHA256SUMS.txt").write_text("\n".join(sorted(checksums)) + "\n", encoding="utf-8")
    print(json.dumps(summary, indent=2))
    return 1 if overall == "FAIL" else 0


if __name__ == "__main__":
    sys.exit(main())
