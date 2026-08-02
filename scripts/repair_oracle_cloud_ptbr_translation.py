#!/usr/bin/env python3
"""Build a source-aligned Oracle Cloud pt-BR DOCX from a machine first pass.

English is authoritative. Exact English matches reuse validated pt-BR
translation memories only when the paragraph contains no hyperlink. Oracle
Cloud-specific terminology is then repaired explicitly.
"""

from __future__ import annotations

import argparse
import copy
import tempfile
import zipfile
from collections import defaultdict
from pathlib import Path

from docx import Document


REPLACEMENTS = {
    "RAG (Resposta a Ações Corretivas e Negativas)": "RAG",
    "MCP (Controle de Pontos de Gerenciamento)": "MCP",
    "TRAPO": "RAG",
    "injeção imediata": "injeção de prompt",
    "Injeção imediata": "Injeção de prompt",
    "Segurança Imediata": "Segurança de Prompts",
    "segurança imediata": "segurança de prompts",
    "processamento imediato": "tratamento de prompts",
    "solicitações ou ferramentas": "prompts ou ferramentas",
    "alerta.": "prompt.",
    "área de aterrissagem": "landing zone",
    "área de destino": "landing zone",
    "locação": "tenancy",
    "seus inquilinos": "suas tenancies",
    "entre locatários": "entre tenancies",
    "posse, domínios de identidade": "tenancy, domínios de identidade",
    "entidades de instância": "instance principals",
    "entidades de recurso": "resource principals",
    "recipientes privilegiados": "contêineres privilegiados",
    "Guarda-corpo": "Guardrail",
    "guarda-corpo": "guardrail",
    "salvaguardas": "guardrails",
    "Incorporação": "Embedding",
    "Loja de vetores": "Repositório vetorial",
    "Candidaturas a Mestrado em Direito": "aplicações de LLM",
    "registros controlados": "registries controlados",
    "Gerenciamento do Ciclo de Vida da IA e do Risco de Modelo": "Gestão do Ciclo de Vida da IA e dos Riscos de Modelo",
    "IA Generativa e Segurança de Prompts": "Inteligência Artificial Generativa e Segurança de Prompts",
    "RAG Segurança": "Segurança de RAG",
    "Laboratório Prático de Caixa de Areia": "Laboratório Prático de Sandbox",
    "procedência": "proveniência",
    "digitalize o conteúdo": "examine o conteúdo",
    "Zonas de Segurança": "Security Zones",
    "Serviço de Varredura de Vulnerabilidades": "Vulnerability Scanning Service",
    "Serviço de Verificação de Vulnerabilidades": "Vulnerability Scanning Service",
    "Hub de Gerenciamento de SO": "OS Management Hub",
    "Hub de Gerenciamento do Sistema Operacional": "OS Management Hub",
    "Cofre OCI, gerenciamento de chaves e segredos": "OCI Vault, Key Management e Secrets",
}


EXACT_REPLACEMENTS = {
    "entidades de recurso do OCI IAM controlam quem e o que pode acessar os recursos. As contas humanas devem usar federação e MFA. As cargas de trabalho devem usar identidades nativas da plataforma e de curta duração em vez de chaves de API incorporadas sempre que possível.":
        "As políticas, os grupos, os grupos dinâmicos, os instance principals e os resource principals do OCI IAM controlam quem e o que pode acessar recursos. Contas humanas devem usar federação e MFA. Sempre que possível, cargas de trabalho devem usar identidades nativas da plataforma e de curta duração, em vez de chaves de API incorporadas.",
    "RAG Security": "Segurança de RAG",
    "Ataques indiretos, codificados, ofuscados e multilíngues de equipe vermelha.":
        "Realize red teaming de ataques multilíngues, codificados, ofuscados e indiretos.",
    "Registro da seleção da ferramenta de log, argumentos, decisão política, resultado e aprovação do usuário.":
        "Registre a seleção da ferramenta, os argumentos, a decisão de política, o resultado e a aprovação do usuário.",
    "Políticas, grupos, grupos dinâmicos, instance principals e resource principals do OCI IAM.":
        "Políticas, grupos, grupos dinâmicos, instance principals e resource principals do OCI IAM",
    "Resistência à injeção, validação de saída, guardrails e aprovação humana.":
        "Resistência à injeção, validação de saída, guardrails e aprovação humana",
    "Lista de permissões de ferramentas, identidade dedicada, aprovação e desativação automática.":
        "Lista de permissões de ferramentas, identidade dedicada, aprovação e mecanismo de desativação imediata",
    "Os serviços de IA adicionam responsabilidades de modelagem de comportamento, tratamento de prompts, avaliação e supervisão humana.":
        "Os serviços de IA acrescentam responsabilidades relacionadas ao comportamento do modelo, ao tratamento de prompts, à avaliação e à supervisão humana.",
    "Considere o uso indevido, a injeção de prompt, a agência excessiva, o envenenamento de dados, o tratamento inseguro da saída e a negação de serviço.":
        "Considere uso indevido, injeção de prompt, autonomia excessiva, envenenamento de dados, tratamento inseguro de saídas e negação de serviço.",
    "Funções separadas para administração, implantação, tempo de execução e auditoria.":
        "Separe as funções de administração, implantação, execução e auditoria.",
    "Atribua a cada agente de IA sua própria identidade e limite o número de ferramentas que podem ser chamadas.":
        "Atribua a cada agente de IA uma identidade própria e limite individualmente cada ferramenta que ele possa chamar.",
    "Registre o fluxo de logs e eventos de borda, e alerte sobre destinos ou volumes incomuns.":
        "Registre eventos de fluxo e de borda e alerte sobre destinos ou volumes incomuns.",
    "Testar o comportamento de DNS, roteamento e failover.":
        "Teste o comportamento de DNS, roteamento e failover.",
    "Coletar somente os dados necessários para a finalidade aprovada.":
        "Colete somente os dados necessários para a finalidade aprovada.",
    "Bloquear segredos e dados regulamentados de solicitações, a menos que sejam explicitamente autorizados.":
        "Bloqueie segredos e dados regulamentados em prompts, salvo quando houver aprovação explícita.",
    "na rede, firewall e exposição pública":
        "Alterações de rede, firewall e exposição pública",
    "Registro de falhas no pipeline e alterações de retenção":
        "Falhas no pipeline de logging e alterações de retenção",
    "Fornecedor do modelo de registro, versão, fontes de dados, licença e limitações conhecidas.":
        "Registre o fornecedor e a versão do modelo, as fontes de dados, a licença e as limitações conhecidas.",
    "Viés de teste, alucinação, injeção de prompt, vazamento de dados e casos de abuso.":
        "Teste vieses, alucinações, injeção de prompt, vazamento de dados e casos de abuso.",
    "Limitar a taxa de operações caras ou sensíveis.":
        "Aplique limites de taxa a operações caras ou sensíveis.",
    "Garanta a ingestão segura, preserve a procedência, digitalize o conteúdo, imponha autorização em nível de documento e exiba citações para que os usuários possam verificar a resposta.":
        "Proteja a ingestão, preserve a proveniência, examine o conteúdo, imponha autorização no nível do documento e exiba citações para que os usuários possam verificar a resposta.",
    "Aprovar e inventariar repositórios de origem.":
        "Aprove e inventarie os repositórios de origem.",
    "Análise da ingestão de dados em busca de malware, segredos, instruções semelhantes a prompts e formatos não suportados.":
        "Examine a ingestão em busca de malware, segredos, instruções semelhantes a prompts e formatos não compatíveis.",
    "Preservar a origem, o proprietário, a classificação, a versão e a hora de ingestão como metadados.":
        "Preserve como metadados a origem, o responsável, a classificação, a versão e o horário de ingestão.",
    "Limitar o tamanho do contexto e impedir a recuperação entre tenancies.":
        "Limite o tamanho do contexto e impeça a recuperação entre tenancies.",
    "Avaliar a precisão da recuperação, a autorização e a exatidão das citações.":
        "Avalie a precisão da recuperação, a autorização e a exatidão das citações.",
    "Inventariar cada agente, ferramenta, servidor MCP, proprietário e fluxo de dados.":
        "Inventarie cada agente, ferramenta, servidor MCP, responsável e fluxo de dados.",
    "Autenticar tanto o cliente quanto o servidor e usar transporte criptografado.":
        "Autentique o cliente e o servidor e use transporte criptografado.",
    "Exigir confirmação para ações irreversíveis, externas ou de alto impacto.":
        "Exija confirmação para ações irreversíveis, externas ou de alto impacto.",
    "Valide os parâmetros e resultados da ferramenta; nunca execute o modelo de texto sem antes verificar.":
        "Valide os parâmetros e as saídas da ferramenta; nunca execute cegamente texto gerado pelo modelo.",
    "Inclui casos normais, extremos, de uso indevido e adversários.":
        "Inclua casos normais, extremos, de uso indevido e adversariais.",
    "Realizar novo teste após alterações no modelo, nos prompts, nos dados, nas ferramentas ou nas políticas.":
        "Teste novamente após alterações no modelo, nos prompts, nos dados, nas ferramentas ou nas políticas.",
    "Contém: revogar credenciais, isolar endpoints, desativar ferramentas, congelar a ingestão ou reverter um modelo.":
        "Conter: revogar credenciais, isolar endpoints, desativar ferramentas, suspender a ingestão ou reverter um modelo.",
    "Analisar custos, exposição pública, identidades, registros e descobertas; exportar um pequeno pacote de evidências.":
        "Analise custos, exposição pública, identidades, logs e constatações; exporte um pequeno pacote de evidências.",
    "Identificação e correção postural":
        "Constatação de postura e correção",
    "relatório de teste de IA": "Relatório de testes de IA",
    "resultados da equipe vermelha com remediação.":
        "Resultados de avaliação e red teaming, com correção",
    "Utilize o treinamento do fornecedor, a documentação oficial e os relatos de boas práticas.":
        "Use treinamentos do fornecedor, documentação oficial e contas seguras para prática.",
    "resource principals do OCI IAM controlam quem e o que pode acessar os recursos. As contas humanas devem usar federação e MFA. As cargas de trabalho devem usar identidades nativas da plataforma e de curta duração em vez de chaves de API incorporadas sempre que possível.":
        "As políticas, os grupos, os grupos dinâmicos, os instance principals e os resource principals do OCI IAM controlam quem e o que pode acessar recursos. Contas humanas devem usar federação e MFA. Sempre que possível, cargas de trabalho devem usar identidades nativas da plataforma e de curta duração, em vez de chaves de API incorporadas.",
    "Auditoria OCI, registro de logs, eventos, notificações, hub do Service Connector e análises de logs.":
        "OCI Audit, Logging, Events, Notifications, Service Connector Hub e Logging Analytics",
    "Prazo": "Termo",
    "Habilite os recursos de auditoria, registro, eventos, notificações, hub do conector de serviço e análise de registro do OCI e direcione eventos administrativos importantes para armazenamento protegido.":
        "Habilite OCI Audit, Logging, Events, Notifications, Service Connector Hub e Logging Analytics e direcione eventos administrativos importantes para um armazenamento protegido.",
    "Configure o OCI Vault, o gerenciamento de chaves e os segredos; armazene um segredo de teste e recupere-o por meio de uma identidade de carga de trabalho, em vez de incorporá-lo ao código.":
        "Configure OCI Vault, Key Management e Secrets; armazene um segredo de teste e recupere-o por meio de uma identidade de carga de trabalho, em vez de incorporá-lo ao código.",
    "Ative ou revise o Cloud Guard, as Security Zones, o Vulnerability Scanning Service e o OS Management Hub; investigue uma das vulnerabilidades encontradas e registre as evidências de correção.":
        "Ative ou revise Cloud Guard, Security Zones, Vulnerability Scanning Service e OS Management Hub; investigue uma constatação e registre as evidências de correção.",
}


def all_paragraphs(document: Document):
    result = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                result.extend(cell.paragraphs)
    return result


def set_text_preserving_first_node(paragraph, value: str) -> None:
    nodes = paragraph._p.xpath(".//w:t")
    if nodes:
        nodes[0].text = value
        for node in nodes[1:]:
            node.text = ""
    else:
        paragraph.add_run(value)


def restore_bookmarks(source, target) -> None:
    starts = source._p.xpath("./w:bookmarkStart")
    ends = source._p.xpath("./w:bookmarkEnd")
    insert_at = 1 if target._p.pPr is not None else 0
    for node in starts:
        target._p.insert(insert_at, copy.deepcopy(node))
        insert_at += 1
    for node in ends:
        target._p.append(copy.deepcopy(node))


def restore_internal_hyperlinks(source, target) -> None:
    source_links = source._p.xpath("./w:hyperlink[@w:anchor]")
    if not source_links:
        return
    if len(source_links) != 1:
        raise SystemExit("unsupported multiple internal hyperlinks in one paragraph")
    hyperlink = copy.deepcopy(source_links[0])
    for child in list(hyperlink):
        hyperlink.remove(child)
    movable = [
        child for child in list(target._p)
        if child.tag.endswith("}r") or child.tag.endswith("}fldSimple")
    ]
    if not movable:
        raise SystemExit("internal hyperlink target has no runs")
    insert_at = 1 if target._p.pPr is not None else 0
    for child in movable:
        target._p.remove(child)
        hyperlink.append(child)
    target._p.insert(insert_at, hyperlink)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--english", type=Path, required=True)
    parser.add_argument("--machine", type=Path, required=True)
    parser.add_argument(
        "--memory", nargs=2, type=Path, action="append", default=[],
        metavar=("ENGLISH", "PORTUGUESE"),
    )
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    english = Document(args.english)
    translated = Document(args.machine)
    source_paragraphs = all_paragraphs(english)
    target_paragraphs = all_paragraphs(translated)

    if len(english.paragraphs) != len(translated.paragraphs):
        raise SystemExit("body paragraph parity failed")
    if len(source_paragraphs) != len(target_paragraphs):
        raise SystemExit("aggregate paragraph parity failed")
    if len(english.tables) != len(translated.tables):
        raise SystemExit("table parity failed")
    if [(len(t.rows), len(t.columns)) for t in english.tables] != [
        (len(t.rows), len(t.columns)) for t in translated.tables
    ]:
        raise SystemExit("table-shape parity failed")
    if len(english.inline_shapes) != len(translated.inline_shapes):
        raise SystemExit("figure parity failed")

    memory = defaultdict(set)
    for memory_english_path, memory_portuguese_path in args.memory:
        memory_source = all_paragraphs(Document(memory_english_path))
        memory_target = all_paragraphs(Document(memory_portuguese_path))
        if len(memory_source) != len(memory_target):
            raise SystemExit(f"translation-memory parity failed: {memory_english_path}")
        for source, target in zip(memory_source, memory_target):
            if source.text.strip():
                memory[source.text].add(target.text)

    # Google Translate turns source bookmarks into internal hyperlinks. Unwrap
    # those synthetic links before restoring the authoritative bookmarks.
    for hyperlink in list(translated._element.xpath(".//w:hyperlink[@w:anchor]")):
        parent = hyperlink.getparent()
        position = parent.index(hyperlink)
        for child in list(hyperlink):
            parent.insert(position, child)
            position += 1
        parent.remove(hyperlink)

    for node in list(translated._element.xpath(".//w:bookmarkStart | .//w:bookmarkEnd")):
        node.getparent().remove(node)

    reused = 0
    for source, target in zip(source_paragraphs, target_paragraphs):
        target_ppr = target._p.get_or_add_pPr()
        for existing in list(target_ppr.xpath("./w:pStyle")):
            target_ppr.remove(existing)
        source_style = source._p.xpath("./w:pPr/w:pStyle")
        if source_style:
            target_ppr.insert(0, copy.deepcopy(source_style[0]))
        restore_internal_hyperlinks(source, target)
        restore_bookmarks(source, target)

        candidates = memory.get(source.text, set())
        if len(candidates) == 1 and not source._p.xpath(".//w:hyperlink") and not target._p.xpath(".//w:hyperlink"):
            set_text_preserving_first_node(target, next(iter(candidates)))
            reused += 1

        if source._p.xpath("./w:hyperlink[@w:anchor]"):
            for node in target._p.xpath(".//w:t"):
                revised_node = node.text or ""
                for old, new in REPLACEMENTS.items():
                    revised_node = revised_node.replace(old, new)
                node.text = EXACT_REPLACEMENTS.get(revised_node, revised_node)
            continue

        revised = target.text
        for old, new in REPLACEMENTS.items():
            revised = revised.replace(old, new)
        revised = EXACT_REPLACEMENTS.get(revised, revised)
        if revised != target.text:
            set_text_preserving_first_node(target, revised)

    starts = translated._element.xpath(".//w:bookmarkStart")
    names = [node.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name") for node in starts]
    if len(names) != len(set(names)):
        raise SystemExit("duplicate bookmark name introduced")
    source_hyperlinks = len(english._element.xpath(".//w:hyperlink"))
    target_hyperlinks = len(translated._element.xpath(".//w:hyperlink"))
    if source_hyperlinks != target_hyperlinks:
        raise SystemExit(f"hyperlink parity failed: {source_hyperlinks} != {target_hyperlinks}")

    translated.core_properties.title = "Manual Prático de Segurança de IA — Oracle Cloud Infrastructure (OCI)"
    translated.core_properties.subject = (
        "Guia prático sobre OCI, identidade, redes, dados, RAG, MCP, agentes, "
        "governança, testes e resposta a incidentes"
    )
    translated.core_properties.keywords = (
        "Oracle Cloud Infrastructure, OCI, IAM, Cloud Guard, RAG, MCP, OWASP, segurança"
    )
    args.output.parent.mkdir(parents=True, exist_ok=True)
    translated.save(args.output)

    # Restore authoritative numbering and styles after the translation service.
    with zipfile.ZipFile(args.english) as source_package:
        numbering = source_package.read("word/numbering.xml")
        styles = source_package.read("word/styles.xml")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temporary:
        temporary_path = Path(temporary.name)
    try:
        with zipfile.ZipFile(args.output) as input_package, zipfile.ZipFile(
            temporary_path, "w", zipfile.ZIP_DEFLATED
        ) as output_package:
            for item in input_package.infolist():
                if item.filename == "word/numbering.xml":
                    payload = numbering
                elif item.filename == "word/styles.xml":
                    payload = styles
                else:
                    payload = input_package.read(item.filename)
                output_package.writestr(item, payload)
        temporary_path.replace(args.output)
    finally:
        temporary_path.unlink(missing_ok=True)

    print(f"translation-memory paragraphs reused: {reused}")


if __name__ == "__main__":
    main()
