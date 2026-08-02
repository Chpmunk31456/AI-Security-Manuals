#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

REPLACEMENTS = {
    "TRAPO": "RAG",
    "injeção imediata": "injeção de prompt",
    "Guarda-corpo": "Guardrail",
    "Incorporação": "Embedding",
    "Prazo": "Termo",
    "um responsável accountable": "um responsável definido",
}


def paragraphs(document: Document):
    result = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                result.extend(cell.paragraphs)
    return result


def set_text(paragraph, value: str) -> None:
    nodes = paragraph._p.xpath(".//w:t")
    if nodes:
        nodes[0].text = value
        for node in nodes[1:]:
            node.text = ""
    else:
        paragraph.add_run(value)


def normalize(value: str) -> str:
    for old, new in REPLACEMENTS.items():
        value = value.replace(old, new)
    return value


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--english", type=Path, required=True)
    parser.add_argument("--custom", type=Path, required=True)
    parser.add_argument("--memory", nargs=2, type=Path, action="append", default=[])
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    custom = {int(k): v for k, v in json.loads(args.custom.read_text(encoding="utf-8")).items()}
    memory: dict[str, str] = {}
    conflicts: set[str] = set()

    for english_path, portuguese_path in args.memory:
        english_paragraphs = paragraphs(Document(english_path))
        portuguese_paragraphs = paragraphs(Document(portuguese_path))
        if len(english_paragraphs) != len(portuguese_paragraphs):
            raise SystemExit(f"paragraph mismatch: {english_path}")
        for source, target in zip(english_paragraphs, portuguese_paragraphs):
            key = source.text.strip()
            value = target.text.strip()
            if not key or not value:
                continue
            if key in memory and memory[key] != value:
                conflicts.add(key)
            else:
                memory[key] = value

    document = Document(args.english)
    target_paragraphs = paragraphs(document)
    missing: list[str] = []

    for index, paragraph in enumerate(target_paragraphs):
        key = paragraph.text.strip()
        if not key:
            continue
        value = custom.get(index) or memory.get(key)
        if not value:
            missing.append(f"{index}: {key}")
            continue
        set_text(paragraph, normalize(value))
        for run in paragraph.runs:
            run._element.get_or_add_rPr().set(qn("w:lang"), "pt-BR")

    if missing:
        raise SystemExit("missing translations:\n" + "\n".join(missing))

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip() == "IBM Cloud | Practical AI Security Manual":
                set_text(paragraph, "IBM Cloud | Manual Prático de Segurança de IA")

    document.core_properties.title = "Manual de Segurança de IA na IBM Cloud"
    document.core_properties.subject = "Segurança, governança e operações de IA na IBM Cloud"
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)

    with zipfile.ZipFile(args.output) as package:
        corrupt = package.testzip()
        if corrupt:
            raise SystemExit(f"corrupt DOCX member: {corrupt}")

    print(
        f"translated={sum(bool(p.text.strip()) for p in target_paragraphs)} "
        f"memory={len(memory)} conflicts={len(conflicts)} output={args.output}"
    )


if __name__ == "__main__":
    main()
