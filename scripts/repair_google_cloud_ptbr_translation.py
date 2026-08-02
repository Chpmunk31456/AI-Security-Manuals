#!/usr/bin/env python3
"""Build a source-aligned Google Cloud pt-BR DOCX from a machine first pass.

English is authoritative. Exact English matches reuse the validated AWS pt-BR
translation memory only when the paragraph contains no hyperlink. Google Cloud
specific terminology is then repaired explicitly.
"""

from __future__ import annotations

import argparse
import copy
import tempfile
import zipfile
from collections import defaultdict
from pathlib import Path

from docx import Document


REPLACEMENTS = {
    "IA de vértices": "Vertex AI",
    "Centro de Comando de Segurança": "Security Command Center",
    "Proteção de Dados Sensíveis": "Sensitive Data Protection",
    "proteção de dados sensíveis": "Sensitive Data Protection",
    "logs de auditoria do Cloud": "Cloud Audit Logs",
    "TRAPO": "RAG",
    "IA Paralela": "shadow AI",
    "IA paralela": "shadow AI",
    "área de aterrissagem": "landing zone",
    "Serviço de Políticas da Organização (OPS)": "Organization Policy Service",
    "NHI (National Health Insurance)": "NHI",
    "Incorporações e similaridade": "Embeddings e similaridade",
    "Estímulo, Contexto, Memória e Segurança do Conteúdo": "Prompt, Contexto, Memória e Segurança de Conteúdo",
    "Agentes governamentais e extensões": "Governar agentes e extensões",
    "Construa uma faixa de pedestres útil.": "Crie um mapeamento cruzado útil.",
    "Use o princípio do menor privilégio e assuma o compromisso.": "Aplique o privilégio mínimo e presuma comprometimento.",
    "Garantia de mudanças de comportamento probabilístico": "O comportamento probabilístico muda a garantia",
    "Modelo de Risco e Resultados": "Modelo de Risco e Constatação",
    "grupos de identidade na nuvem": "grupos do Cloud Identity",
    "Federação de Identidade de Carga de Trabalho": "Workload Identity Federation",
    "Gerenciador de Acesso Privilegiado": "Privileged Access Manager",
    "Catálogo Universal Dataplex": "Dataplex Universal Catalog",
    "análises de integridade de segurança": "Security Health Analytics",
    "detecção de ameaças por eventos": "Event Threat Detection",
    "logs de fluxo da VPC": "VPC Flow Logs",
    "Gemini para Google Workspace": "Gemini for Google Workspace",
    "contas de serviço do IAM": "IAM service accounts",
}


def all_paragraphs(document: Document):
    result = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                result.extend(cell.paragraphs)
    return result


def set_text_preserving_first_node(paragraph, value: str) -> None:
    nodes = paragraph._p.xpath(".//w:t")
    if nodes:
        nodes[0].text = value
        for node in nodes[1:]:
            node.text = ""
    else:
        paragraph.add_run(value)


def restore_bookmarks(source, target) -> None:
    starts = source._p.xpath("./w:bookmarkStart")
    ends = source._p.xpath("./w:bookmarkEnd")
    insert_at = 1 if target._p.pPr is not None else 0
    for node in starts:
        target._p.insert(insert_at, copy.deepcopy(node))
        insert_at += 1
    for node in ends:
        target._p.append(copy.deepcopy(node))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--english", type=Path, required=True)
    parser.add_argument("--machine", type=Path, required=True)
    parser.add_argument("--memory-english", type=Path, required=True)
    parser.add_argument("--memory-portuguese", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    english = Document(args.english)
    translated = Document(args.machine)
    memory_english = Document(args.memory_english)
    memory_portuguese = Document(args.memory_portuguese)
    source_paragraphs = all_paragraphs(english)
    target_paragraphs = all_paragraphs(translated)
    memory_source = all_paragraphs(memory_english)
    memory_target = all_paragraphs(memory_portuguese)

    if len(english.paragraphs) != len(translated.paragraphs):
        raise SystemExit("body paragraph parity failed")
    if len(source_paragraphs) != len(target_paragraphs):
        raise SystemExit("aggregate paragraph parity failed")
    if len(english.tables) != len(translated.tables):
        raise SystemExit("table parity failed")
    if [(len(t.rows), len(t.columns)) for t in english.tables] != [
        (len(t.rows), len(t.columns)) for t in translated.tables
    ]:
        raise SystemExit("table-shape parity failed")
    if len(english.inline_shapes) != len(translated.inline_shapes):
        raise SystemExit("figure parity failed")

    memory = defaultdict(set)
    for source, target in zip(memory_source, memory_target):
        if source.text.strip():
            memory[source.text].add(target.text)

    for node in list(translated._element.xpath(".//w:bookmarkStart | .//w:bookmarkEnd")):
        node.getparent().remove(node)

    reused = 0
    for source, target in zip(source_paragraphs, target_paragraphs):
        target_ppr = target._p.get_or_add_pPr()
        for existing in list(target_ppr.xpath("./w:pStyle")):
            target_ppr.remove(existing)
        source_style = source._p.xpath("./w:pPr/w:pStyle")
        if source_style:
            target_ppr.insert(0, copy.deepcopy(source_style[0]))
        restore_bookmarks(source, target)

        candidates = memory.get(source.text, set())
        if len(candidates) == 1 and not source._p.xpath(".//w:hyperlink") and not target._p.xpath(".//w:hyperlink"):
            set_text_preserving_first_node(target, next(iter(candidates)))
            reused += 1

        revised = target.text
        for old, new in REPLACEMENTS.items():
            revised = revised.replace(old, new)
        if revised != target.text:
            set_text_preserving_first_node(target, revised)

    starts = translated._element.xpath(".//w:bookmarkStart")
    names = [node.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name") for node in starts]
    if len(names) != len(set(names)):
        raise SystemExit("duplicate bookmark name introduced")

    translated.core_properties.title = "Proteção de IA Empresarial no Google Cloud"
    translated.core_properties.subject = (
        "Guia prático sobre Google Cloud, Vertex AI, Gemini, identidade, dados, RAG, "
        "MCP, agentes, governança e resposta a incidentes"
    )
    translated.core_properties.keywords = (
        "Google Cloud, Vertex AI, Gemini, IAM, Dataplex, RAG, MCP, OWASP, segurança"
    )
    args.output.parent.mkdir(parents=True, exist_ok=True)
    translated.save(args.output)

    with zipfile.ZipFile(args.english) as source_package:
        numbering = source_package.read("word/numbering.xml")
        styles = source_package.read("word/styles.xml")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temporary:
        temporary_path = Path(temporary.name)
    try:
        with zipfile.ZipFile(args.output) as input_package, zipfile.ZipFile(
            temporary_path, "w", zipfile.ZIP_DEFLATED
        ) as output_package:
            for item in input_package.infolist():
                if item.filename == "word/numbering.xml":
                    payload = numbering
                elif item.filename == "word/styles.xml":
                    payload = styles
                else:
                    payload = input_package.read(item.filename)
                output_package.writestr(item, payload)
        temporary_path.replace(args.output)
    finally:
        temporary_path.unlink(missing_ok=True)

    print(f"translation-memory paragraphs reused: {reused}")


if __name__ == "__main__":
    main()
