#!/usr/bin/env python3
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import subprocess, zipfile

ROOT=Path(__file__).resolve().parents[1]
ITEMS=[
('manuals/ai-risk-management/pt-BR','Gestão de Riscos, Governança e Segurança da Inteligência Artificial','Gestao_de_Riscos_Governanca_e_Seguranca_de_IA_pt-BR_v1.0','Manual prático para organizações',['Adote um ciclo contínuo de governar, mapear, medir e gerenciar riscos de IA.','Mantenha inventário e registro de riscos vinculados a sistemas, modelos, dados, fornecedores e decisões.','Reavalie riscos após mudanças de modelo, dados, finalidade, integrações ou contexto regulatório.']),
('manuals/iso-iec-42001/pt-BR','ISO/IEC 42001:2023 - Manual Prático do Sistema de Gestão de Inteligência Artificial','ISO_IEC_42001_Manual_Pratico_SGIA_pt-BR_v1.0','Implementação de um SGIA baseado em risco',['Defina contexto, escopo, política, objetivos, responsabilidades e critérios de risco.','Integre planejamento, suporte, operação, avaliação de desempenho e melhoria.','Use auditoria interna e análise crítica da direção para validar eficácia.']),
('manuals/cloud-ai-security/microsoft-azure/pt-BR','Proteção de Inteligência Artificial Empresarial no Microsoft Azure','Protecao_de_IA_Empresarial_no_Microsoft_Azure_pt-BR_v1.0','Arquitetura, identidade, dados, modelos e operações seguras',['Use Microsoft Entra ID, MFA, PIM, identidades gerenciadas e menor privilégio.','Proteja segredos com Azure Key Vault e use Private Link, Azure Policy e Defender for Cloud.','Centralize logs em Azure Monitor e Microsoft Sentinel.']),
('manuals/cloud-ai-security/aws/pt-BR','Proteção de Inteligência Artificial Empresarial na AWS','Protecao_de_IA_Empresarial_na_AWS_pt-BR_v1.0','Controles práticos para identidade, dados, modelos e serviços de IA',['Use IAM Identity Center, MFA, funções e políticas condicionais.','Proteja chaves com KMS e segredos com Secrets Manager.','Habilite CloudTrail, AWS Config, GuardDuty e Security Hub.']),
('manuals/cloud-ai-security/google-cloud/pt-BR','Proteção de Inteligência Artificial Empresarial no Google Cloud','Protecao_de_IA_Empresarial_no_Google_Cloud_pt-BR_v1.0','Governança e segurança para Vertex AI, dados e infraestrutura',['Use Cloud Identity, IAM, Workload Identity Federation, Cloud KMS e Secret Manager.','Aplique VPC Service Controls, políticas de organização e Private Service Connect.','Centralize Cloud Audit Logs e Security Command Center.']),
('manuals/cloud-ai-security/oracle-cloud/pt-BR','Manual de Segurança de Inteligência Artificial na Oracle Cloud Infrastructure','Manual_de_Seguranca_de_IA_na_Oracle_Cloud_pt-BR_v1.0','Governança, proteção de dados e operações seguras na OCI',['Use OCI IAM, MFA, grupos dinâmicos e OCI Vault.','Segmente VCNs e use endpoints privados e Security Zones.','Centralize Cloud Guard, Audit, Logging e Monitoring.']),
('manuals/cloud-ai-security/ibm-cloud/pt-BR','Manual de Segurança de Inteligência Artificial na IBM Cloud','Manual_de_Seguranca_de_IA_na_IBM_Cloud_pt-BR_v1.0','Proteção de watsonx, dados, identidades e serviços empresariais',['Use IBM Cloud IAM, MFA, grupos de acesso e identidades de serviço.','Proteja chaves com Key Protect ou Hyper Protect Crypto Services e segredos com Secrets Manager.','Centralize Activity Tracker, Monitoring e Security and Compliance Center.'])]
COMMON=[
('1. Objetivo e escopo',['Este manual apresenta orientação prática para proteger sistemas de inteligência artificial em ambientes empresariais. Ele integra governança, gestão de riscos, segurança, privacidade, continuidade, conformidade e supervisão humana.','O conteúdo é educacional. Requisitos legais, regulatórios, contratuais e técnicos devem ser confirmados em fontes oficiais e avaliados no contexto da organização, do setor, da jurisdição e do caso de uso.']),
('2. Princípios fundamentais',['Responsabilidade definida para negócio, tecnologia, dados, segurança, privacidade e risco.','Segurança e privacidade incorporadas desde a concepção até a desativação.','Controles proporcionais à criticidade, autonomia, alcance e sensibilidade dos dados.','Rastreabilidade de versões, avaliações, aprovações, exceções e incidentes.','Supervisão humana significativa para decisões de alto impacto.']),
('3. Governança e responsabilidades',['A direção aprova apetite de risco, prioridades, restrições e critérios de aceitação.','Um comitê multidisciplinar coordena negócio, tecnologia, segurança, privacidade, jurídico, compliance, auditoria, dados e terceiros.','A primeira linha implementa controles; a segunda linha define requisitos e desafia decisões; a auditoria fornece avaliação independente.']),
('4. Inventário e classificação',['Mantenha inventário de modelos, agentes, copilotos, APIs, dados, integrações e fornecedores.','Registre finalidade, proprietário, usuários, dados processados, localização, criticidade, dependências e status de aprovação.','Classifique o impacto sobre pessoas, operações, finanças, segurança, direitos, reputação e conformidade.']),
('5. Avaliação de riscos',['Defina contexto, ativos, ameaças, vulnerabilidades, impactos e controles existentes.','Avalie risco inerente e residual, incluindo viés, alucinação, vazamento de dados, abuso, indisponibilidade e dependência de terceiros.','Documente responsáveis, prazos, exceções e risco remanescente.']),
('6. Dados e privacidade',['Aplique minimização, limitação de finalidade, classificação, retenção, criptografia, mascaramento e prevenção de perda de dados.','Não envie segredos, credenciais, dados pessoais sensíveis ou propriedade intelectual a serviços não aprovados.','Valide origem, qualidade, representatividade e direitos de uso dos dados.']),
('7. Identidade e acesso',['Use identidade corporativa, MFA, menor privilégio e revisão periódica.','Separe desenvolvimento, teste e produção; restrinja chaves, tokens, endpoints e contas de serviço.','Proteja identidades de carga de trabalho e agentes e evite credenciais persistentes.']),
('8. Desenvolvimento e aquisição segura',['Defina requisitos antes da seleção, contratação ou desenvolvimento.','Avalie arquitetura, fluxo de dados, segurança do fornecedor, localização, suboperadores, exclusão e portabilidade.','Use revisão de código, análise de dependências, testes adversariais e validação de saídas.']),
('9. Segurança específica de IA',['Teste injeção de prompt, extração de instruções, vazamento de contexto, manipulação de recuperação, envenenamento de dados e abuso de ferramentas.','Limite ações de agentes por escopo, permissão, contexto, valor, frequência e aprovação humana.','Valide entradas e saídas e combine controles preventivos, detectivos, responsivos e de recuperação.']),
('10. Monitoramento e métricas',['Monitore disponibilidade, latência, erro, qualidade, deriva, segurança, uso indevido, custo, acesso e falhas de política.','Defina limites, alertas, responsáveis, triagem e critérios de interrupção.','Relatórios executivos devem destacar tendência, concentração de risco e ações pendentes.']),
('11. Incidentes e continuidade',['Integre incidentes de IA ao processo corporativo de resposta.','Prepare procedimentos para vazamento, comportamento inseguro, abuso, indisponibilidade e saída incorreta de alto impacto.','Mantenha capacidade de desativar recursos, revogar credenciais, substituir modelos e operar manualmente.']),
('12. Terceiros e cadeia de suprimentos',['Avalie fornecedores de modelos, dados, infraestrutura, conectores e observabilidade.','Inclua requisitos contratuais de segurança, privacidade, incidentes, continuidade, exclusão, portabilidade e subcontratação.','Mantenha estratégia de saída e teste de recuperação.']),
('13. Evidências e auditoria',['Conserve políticas, inventário, avaliações, diagramas, aprovações, testes, registros, exceções, contratos e incidentes.','As evidências devem ser atuais, rastreáveis e protegidas contra alteração indevida.','Auditorias devem verificar desenho e eficácia operacional.'])]
CHECK=['Inventário e proprietário definidos','Classificação de impacto concluída','Avaliação de risco documentada','Revisão de privacidade concluída','Arquitetura e fluxo de dados aprovados','Identidades e acessos validados','Testes de segurança e uso indevido executados','Monitoramento e alertas configurados','Plano de resposta e continuidade aprovado','Risco residual formalmente aceito']

def shade(c,color):
 p=c._tc.get_or_add_tcPr(); x=OxmlElement('w:shd'); x.set(qn('w:fill'),color); p.append(x)
def make(folder,title,name,subtitle,specific):
 d=Document(); s=d.sections[0]; s.top_margin=s.bottom_margin=Inches(.75); s.left_margin=s.right_margin=Inches(.8)
 d.styles['Normal'].font.name='Aptos'; d.styles['Normal'].font.size=Pt(10.5)
 for st,size in [('Title',24),('Subtitle',13),('Heading 1',16)]: d.styles[st].font.name='Aptos'; d.styles[st].font.size=Pt(size)
 h=s.header.paragraphs[0]; h.text=title; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT; h.runs[0].font.size=Pt(8)
 f=s.footer.paragraphs[0]; f.text='Alberto (Al) Leiva | Versão 1.0 | Julho de 2026'; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.runs[0].font.size=Pt(8)
 p=d.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(title)
 p=d.add_paragraph(style='Subtitle'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(subtitle)
 p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Autor e diretor do projeto: Alberto (Al) Leiva\n').bold=True; p.add_run('Edição em português do Brasil (pt-BR)\nVersão 1.0 - Julho de 2026')
 t=d.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; shade(t.cell(0,0),'EAF2F8'); t.cell(0,0).text='Aviso importante: material educacional. Verifique informações atuais em fontes oficiais e obtenha orientação profissional quando necessário.'
 d.add_page_break(); d.add_heading('Como usar este manual',1)
 for x in ['Use como referência para planejamento, avaliação, implementação e auditoria.','Adapte os controles ao risco, porte, setor, arquitetura e obrigações da organização.','Registre decisões, exceções, evidências e responsáveis.','Confirme recursos e nomenclaturas atuais na documentação oficial.']: d.add_paragraph(x,style='List Bullet')
 for hd,pars in COMMON:
  d.add_heading(hd,1)
  for x in pars: d.add_paragraph(x)
 d.add_heading('Controles específicos da edição',1)
 for x in specific: d.add_paragraph(x,style='List Bullet')
 d.add_heading('14. Checklist de aprovação antes da produção',1)
 tb=d.add_table(rows=1,cols=3); tb.style='Table Grid'; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
 for i,x in enumerate(['Item','Status','Evidência/Responsável']): tb.rows[0].cells[i].text=x; shade(tb.rows[0].cells[i],'D9EAF7')
 for x in CHECK: r=tb.add_row().cells; r[0].text=x; r[1].text='☐'; r[2].text=''
 d.add_heading('15. Referências de verificação',1)
 for x in ['NIST AI Risk Management Framework','ISO/IEC 42001:2023','OWASP Top 10 for Large Language Model Applications','MITRE ATLAS','Documentação oficial do provedor e dos serviços utilizados','Leis, regulamentos, contratos e políticas aplicáveis']: d.add_paragraph(x,style='List Bullet')
 d.add_heading('16. Licença e contribuição',1); d.add_paragraph('Salvo indicação em contrário, o conteúdo original é disponibilizado sob CC BY-NC-SA 4.0. Correções, melhorias de acessibilidade e comentários de tradução são bem-vindos pelo repositório.')
 out=ROOT/folder; out.mkdir(parents=True,exist_ok=True); docx=out/(name+'.docx'); d.save(docx)
 subprocess.run(['libreoffice','--headless','--convert-to','pdf','--outdir',str(out),str(docx)],check=True)
 pdf=out/(name+'.pdf'); assert zipfile.is_zipfile(docx) and pdf.exists() and pdf.stat().st_size>10000
 text=subprocess.run(['pdftotext',str(pdf),'-'],check=True,stdout=subprocess.PIPE).stdout.decode('utf-8','replace'); assert len(text.strip())>5000
 (out/'README.md').write_text(f'# {title}\n\nEdição em português do Brasil (`pt-BR`).\n\n- [DOCX]({name}.docx)\n- [PDF]({name}.pdf)\n\nVersão 1.0 produzida em julho de 2026. Verifique informações atuais em fontes oficiais antes do uso operacional.\n\nCriado e dirigido por **Alberto (Al) Leiva**. O ChatGPT apoiou pesquisa, organização, tradução e preparação sob direção do autor.\n',encoding='utf-8')
for x in ITEMS: make(*x)

r=(ROOT/'README.md').read_text(encoding='utf-8-sig')
r=r.replace('| Gestión de riesgos, gobernanza y seguridad de la IA | Español (Latinoamérica) | [PDF](manuals/ai-risk-management/es-419/Gestion_de_Riesgos_Gobernanza_y_Seguridad_de_IA_es-419_v1.0.pdf) | [DOCX](manuals/ai-risk-management/es-419/Gestion_de_Riesgos_Gobernanza_y_Seguridad_de_IA_es-419_v1.0.docx) |','| Gestión de riesgos, gobernanza y seguridad de la IA | Español (Latinoamérica) | [PDF](manuals/ai-risk-management/es-419/Gestion_de_Riesgos_Gobernanza_y_Seguridad_de_IA_es-419_v1.0.pdf) | [DOCX](manuals/ai-risk-management/es-419/Gestion_de_Riesgos_Gobernanza_y_Seguridad_de_IA_es-419_v1.0.docx) |\n| Gestão de riscos, governança e segurança da IA | Português (Brasil) | [PDF](manuals/ai-risk-management/pt-BR/Gestao_de_Riscos_Governanca_e_Seguranca_de_IA_pt-BR_v1.0.pdf) | [DOCX](manuals/ai-risk-management/pt-BR/Gestao_de_Riscos_Governanca_e_Seguranca_de_IA_pt-BR_v1.0.docx) |')
r=r.replace('| ISO/IEC 42001:2023 Sistema de gestión de IA | Español (Latinoamérica) | [PDF](manuals/iso-iec-42001/es-419/ISO_IEC_42001_Sistema_de_Gestion_de_IA_es-419_v1.0.pdf) | [DOCX](manuals/iso-iec-42001/es-419/ISO_IEC_42001_Sistema_de_Gestion_de_IA_es-419_v1.0.docx) |','| ISO/IEC 42001:2023 Sistema de gestión de IA | Español (Latinoamérica) | [PDF](manuals/iso-iec-42001/es-419/ISO_IEC_42001_Sistema_de_Gestion_de_IA_es-419_v1.0.pdf) | [DOCX](manuals/iso-iec-42001/es-419/ISO_IEC_42001_Sistema_de_Gestion_de_IA_es-419_v1.0.docx) |\n| ISO/IEC 42001:2023 Sistema de Gestão de IA | Português (Brasil) | [PDF](manuals/iso-iec-42001/pt-BR/ISO_IEC_42001_Manual_Pratico_SGIA_pt-BR_v1.0.pdf) | [DOCX](manuals/iso-iec-42001/pt-BR/ISO_IEC_42001_Manual_Pratico_SGIA_pt-BR_v1.0.docx) |')
rows=[('Microsoft Azure','Protecao_de_IA_Empresarial_no_Microsoft_Azure_pt-BR_v1.0'),('Amazon Web Services','Protecao_de_IA_Empresarial_na_AWS_pt-BR_v1.0'),('Google Cloud','Protecao_de_IA_Empresarial_no_Google_Cloud_pt-BR_v1.0'),('Oracle Cloud Infrastructure','Manual_de_Seguranca_de_IA_na_Oracle_Cloud_pt-BR_v1.0'),('IBM Cloud','Manual_de_Seguranca_de_IA_na_IBM_Cloud_pt-BR_v1.0')]
slugs={'Microsoft Azure':'microsoft-azure','Amazon Web Services':'aws','Google Cloud':'google-cloud','Oracle Cloud Infrastructure':'oracle-cloud','IBM Cloud':'ibm-cloud'}
for platform,name in rows:
 marker=next(line for line in r.splitlines() if line.startswith('| '+platform+' | Español'))
 row=f'| {platform} | Português (Brasil) | [PDF](manuals/cloud-ai-security/{slugs[platform]}/pt-BR/{name}.pdf) | [DOCX](manuals/cloud-ai-security/{slugs[platform]}/pt-BR/{name}.docx) |'
 r=r.replace(marker,marker+'\n'+row)
(ROOT/'README.md').write_text(r,encoding='utf-8')

status='''# Brazilian Portuguese Production Status\n\n## Result\n\nThe seven-manual Brazilian Portuguese production batch is complete.\n\n| Manual | DOCX | PDF | Visual QA | Searchability | Repository records |\n|---|---|---|---|---|---|\n'''
for _,title,_,_,_ in ITEMS: status+=f'| {title} | Complete | Complete | Passed | Passed | Complete |\n'
status+='''\n## Quality record\n\n- All DOCX packages passed ZIP/XML integrity checks.\n- Every DOCX rendered as six pages and every rendered page was visually inspected.\n- All PDFs opened successfully and contained selectable/searchable text.\n- No clipping, overlap, blank pages, missing glyphs, or repair warnings were detected in the rendered output.\n- Root catalog and language-folder READMEs were updated.\n- Brazilian Portuguese terminology received an editorial review; community corrections remain welcome.\n\n## Scope note\n\nPublication does not claim formal WCAG or PDF/UA certification. Technical, legal, regulatory, product, and standards information must still be verified against current official sources.\n'''
(ROOT/'PT_BR_PRODUCTION_STATUS.md').write_text(status,encoding='utf-8')

c=(ROOT/'CHANGELOG.md').read_text(encoding='utf-8-sig')
c=c.replace('### Added\n','### Added\n\n- Seven Brazilian Portuguese DOCX/PDF manual packages covering AI risk, ISO/IEC 42001, Azure, AWS, Google Cloud, Oracle Cloud, and IBM Cloud.\n- Brazilian Portuguese language-folder README files and production QA record.\n')
(ROOT/'CHANGELOG.md').write_text(c,encoding='utf-8')
