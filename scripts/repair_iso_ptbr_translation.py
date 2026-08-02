#!/usr/bin/env python3
"""Repair protected ISO/IEC terminology in the machine-assisted pt-BR DOCX.

The script is deliberately scoped to the ISO/IEC 42001 manual and only fixes
reproducible translation defects where product names, the AIMS acronym, or
section identifiers were translated or localized. It does not add content.
"""

from pathlib import Path
import re

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
ENGLISH = ROOT / "manuals/iso-iec-42001/en/ISO_IEC_42001_Practical_AIMS_Manual_en_v1.0.docx"
PORTUGUESE = ROOT / "manuals/iso-iec-42001/pt-BR/ISO_IEC_42001_Manual_Pratico_SGIA_pt-BR_v1.0.docx"

PROTECTED_REPLACEMENTS = {
    "OpenMetadados": "OpenMetadata",
    "Grandes Expectativas": "Great Expectations",
    "Evidentemente": "Evidently",
    "Verificações profundas": "Deepchecks",
    "Inspecionar IA": "Inspect AI",
    "Presídio": "Presidio",
    "Agente de Política Aberta": "Open Policy Agent",
    "AIMS (Sistemas de Gestão de Ações Integradas)": "AIMS",
    "AIMS (Acordo de Incentivo à Implementação de Medidas de Segurança)": "AIMS",
    "AIMS (Sistemas de Gestão de Inteligência Artificial)": "AIMS",
    "AIMS (Sistema de Gestão de Ações Integradas)": "AIMS",
    "RAG (Raiz, Verde e Gramática)": "RAG",
    "estado de ação (SoA)": "Declaração de Aplicabilidade (SoA)",
    "Declaração de Ações (SoA)": "Declaração de Aplicabilidade (SoA)",
    "Declaração de Acordo (SoA)": "Declaração de Aplicabilidade (SoA)",
    "SoA (Solução de Acordo)": "SoA",
    "SoA (Solução de Ação)": "SoA",
    "SoA (Situação das Atividades)": "SoA",
}


def replace_text_nodes(paragraph, replacements):
    for node in paragraph._element.xpath(".//w:t"):
        text = node.text or ""
        for old, new in replacements.items():
            text = text.replace(old, new)
        node.text = text


def main():
    en = Document(ENGLISH)
    pt = Document(PORTUGUESE)
    if len(en.paragraphs) != len(pt.paragraphs) or len(en.tables) != len(pt.tables):
        raise SystemExit("Structural alignment failed; refusing to patch")

    for index, (source, target) in enumerate(zip(en.paragraphs, pt.paragraphs)):
        replace_text_nodes(target, PROTECTED_REPLACEMENTS)

        # Google localized some clause/section identifiers as decimals.
        source_prefix = re.match(r"^(\d+(?:\.\d+)*)\b", source.text)
        target_prefix = re.match(r"^(\d+(?:[.,]\d+)*)\b", target.text)
        if source_prefix and target_prefix and source_prefix.group(1) != target_prefix.group(1):
            replace_text_nodes(target, {target_prefix.group(1): source_prefix.group(1)})

        # AIMS is a defined ISO/IEC 42001 concept and must not become a verb,
        # a generic objective, or a translated ordinary noun.
        if "AIMS" in source.text and "AIMS" not in target.text:
            contextual = {
                109: {"Ferramentas de IA de código aberto para evidências e garantia de IA": "Ferramentas de código aberto para evidências do AIMS e garantia de IA"},
                191: {"OBJETIVOS": "AIMS"},
                203: {"objetivos": "AIMS"},
            }
            if index not in contextual:
                raise SystemExit(f"Unreviewed AIMS loss in paragraph {index}: {target.text}")
            replace_text_nodes(target, contextual[index])

        if "SoA" in source.text and "SoA" not in target.text:
            contextual_soa = {
                357: {"alterações no Estado da Arte": "alterações na SoA"},
                384: {"Declaração de Acordo": "Declaração de Aplicabilidade (SoA)"},
            }
            if index not in contextual_soa:
                raise SystemExit(f"Unreviewed SoA loss in paragraph {index}: {target.text}")
            replace_text_nodes(target, contextual_soa[index])

    for table_index, (source_table, target_table) in enumerate(zip(en.tables, pt.tables)):
        source_cells = [cell for row in source_table.rows for cell in row.cells]
        target_cells = [cell for row in target_table.rows for cell in row.cells]
        if len(source_cells) != len(target_cells):
            raise SystemExit(f"Cell-count mismatch in table {table_index}")
        for cell_index, (source_cell, target_cell) in enumerate(zip(source_cells, target_cells)):
            for paragraph in target_cell.paragraphs:
                replace_text_nodes(paragraph, PROTECTED_REPLACEMENTS)
            if "AIMS" in source_cell.text and "AIMS" not in target_cell.text:
                reviewed = {(3, 3), (12, 4), (38, 2)}
                if (table_index, cell_index) not in reviewed:
                    raise SystemExit(
                        f"Unreviewed AIMS loss in table {table_index}, cell {cell_index}: {target_cell.text}"
                    )
                for paragraph in target_cell.paragraphs:
                    replace_text_nodes(paragraph, {"MIRA": "AIMS", "OBJETIVOS": "AIMS"})
            if "SoA" in source_cell.text and "SoA" not in target_cell.text:
                soa_cells = {
                    (9, 0): {"Aviso da Declaração de Aplicabilidade:": "Aviso da Declaração de Aplicabilidade (SoA):"},
                    (14, 5): {"estado de ação": "SoA"},
                    (37, 8): {"nível de atividade": "SoA"},
                    (39, 32): {"nível de atividade": "SoA"},
                }
                if (table_index, cell_index) not in soa_cells:
                    raise SystemExit(
                        f"Unreviewed SoA loss in table {table_index}, cell {cell_index}: {target_cell.text}"
                    )
                for paragraph in target_cell.paragraphs:
                    replace_text_nodes(paragraph, soa_cells[(table_index, cell_index)])

    pt.save(PORTUGUESE)


if __name__ == "__main__":
    main()
