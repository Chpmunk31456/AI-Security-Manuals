#!/usr/bin/env python3
"""Repair reproducible Azure pt-BR machine-translation defects.

English remains the authority.  The input must preserve the English DOCX's
paragraph/table structure; this script fails closed if structural alignment is
lost.  Corrections below are limited to protected product names, acronyms,
security terminology, and documented mistranslations found during the scoped
Microsoft Azure reconstruction.
"""

from __future__ import annotations

import argparse
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


GLOBAL_REPLACEMENTS = {
    "Âmbito de aplicação": "Purview",
    "Defender para Nuvem": "Defender for Cloud",
    "Política do Azure": "Azure Policy",
    "Gerenciamento de API": "API Management",
    "Cofre de chaves": "Key Vault",
    "Confiança Zero": "Zero Trust",
    "API da OpenAI": "OpenAI API",
    "API OpenAI": "OpenAI API",
    "assuma o compromisso": "presuma comprometimento",
    "injeção imediata": "injeção de prompt",
    "injeção indireta por estímulo": "injeção indireta de prompt",
    "vazamento imediato do sistema": "vazamento do prompt do sistema",
    "equipe vermelha": "red teaming",
    "Equipe vermelha": "Red teaming",
    "treinamento de equipes vermelhas": "red teaming",
    "teste de intrusão com IA": "red teaming de IA",
}


# Aggregate paragraph indexes include body paragraphs followed by every table
# cell paragraph, in document order.  The alignment checks below make these
# indexes safe and reproducible for this source version.
OVERRIDES = {
    96: "Capítulo 11 — Prompt, Contexto, Memória e Segurança de Conteúdo\t40",
    363: (
        "Um embedding representa o conteúdo como números para que um sistema possa comparar significado "
        "ou similaridade. Embeddings dão suporte à pesquisa semântica e ao RAG. Eles são dados derivados e "
        "ainda podem revelar informações sobre a fonte. Um banco de dados vetorial armazena e pesquisa essas "
        "representações juntamente com identificadores e metadados."
    ),
    384: "Treinamento, inferência, ajuste fino e RAG são conceitos distintos.",
    443: "Prompts do sistema, templates e estado da conversa",
    536: "Declare quais controles continuam eficazes se as defesas contra ataques por prompt falharem.",
    550: "Prompts do sistema, políticas e conhecimento proprietário",
    563: "Injeção de prompt, envenenamento, extração, evasão e exfiltração",
    603: "Versione prompts, código, implantações de modelos, conjuntos de dados e configurações de avaliação",
    645: "Compare a OpenAI API direta, uma rota OpenAI hospedada no Azure, a API direta do Claude e o Claude por meio do Microsoft Foundry.",
    655: "Capítulo 11 — Prompt, Contexto, Memória e Segurança de Conteúdo",
    656: (
        "Trate prompts e contexto como componentes controlados da aplicação, reconhecendo que eles não são "
        "limites de segurança."
    ),
    658: (
        "Escreva instruções claras que definam função, tarefas permitidas, ações proibidas, escalonamento, "
        "uso de fontes, incerteza, citações e comportamento das ferramentas. Armazene e versione prompts como "
        "código. Um prompt do sistema pode orientar o comportamento, mas usuários ou conteúdo recuperado ainda "
        "podem manipulá-lo."
    ),
    713: "Detecte injeção indireta de prompt e instruções suspeitas",
    833: (
        "O OWASP Non-Human Identities Top 10 destaca desligamento inadequado, vazamento de segredos, "
        "identidades vulneráveis de terceiros, autenticação insegura, privilégio excessivo, configuração de "
        "nuvem insegura, segredos de longa duração, isolamento fraco de ambientes, reutilização de identidade e "
        "uso humano de NHIs. Trate essa lista como um checklist de projeto e auditoria."
    ),
    847: "Prática: crie um cadastro de NHI",
    854: "Cada NHI é única, tem proprietário e finalidade definidos e está inventariada.",
    885: (
        "Defina finalidade, base legal, aviso, consentimento quando exigido, minimização, acesso, correção, "
        "retenção, exclusão, transferência internacional e tratamento de direitos antes de processar "
        "informações pessoais. Prompts, saídas, embeddings, logs e revisões humanas também podem conter dados "
        "pessoais."
    ),
    951: (
        "A lista de 2025 para aplicações LLM abrange injeção de prompt, divulgação de informações sensíveis, "
        "risco da cadeia de suprimentos, envenenamento de dados e modelos, tratamento inadequado de saídas, "
        "autonomia excessiva, vazamento do prompt do sistema, fragilidades de vetores e embeddings, "
        "desinformação e consumo sem limites. Mapeie cada risco à arquitetura e aos testes reais."
    ),
    952: "A injeção de prompt exige separação de entradas, privilégio mínimo, aprovação e monitoramento",
    1050: "Teste injeção direta e indireta de prompt",
    1088: (
        "Um rastreamento útil conecta a identidade humana ou da carga de trabalho, a sessão, o agente, a "
        "implantação do modelo, a versão do template de prompt, as fontes de dados, os identificadores dos "
        "documentos recuperados, os filtros, as chamadas de ferramentas, as decisões de autorização, os "
        "eventos de segurança de conteúdo, a resposta e o custo. Evite registrar segredos ou conteúdo sensível "
        "desnecessário."
    ),
    1102: (
        "Os playbooks devem abranger divulgação sensível, NHI comprometida, injeção de prompt, RAG envenenado, "
        "MCP ou conector malicioso, saída prejudicial, ação insegura, comprometimento do provedor, indisponibilidade "
        "e custo descontrolado. Defina contenção que não destrua evidências."
    ),
    1106: "Reverta o modelo, o prompt, o código ou a política",
    1235: (
        "Prompts, respostas, aprovações, versões de modelos, avaliações e ações de agentes podem se tornar "
        "registros comerciais ou regulatórios. Coordene retenção, preservação legal, eDiscovery, privacidade, "
        "investigação e minimização. Manter tudo para sempre cria risco; excluir tudo imediatamente destrói a "
        "responsabilização."
    ),
    1315: (
        "Inventarie IA desenvolvida internamente, adquirida, incorporada, experimental e shadow AI. Registre "
        "finalidade, proprietário, usuários, pessoas afetadas, modelo, provedor, dados, agentes, NHIs, integrações, "
        "região, nível de risco, aprovações, monitoramento, incidentes, data de revisão e status de desativação."
    ),
    1380: "Exercite injeção direta e indireta de prompt",
    1399: (
        "Reúna arquitetura, identidade, RAG, controles de agentes, testes, governança e operações em um único "
        "projeto baseado em evidências."
    ),
    1418: "Revise NHIs e segredos",
    1456: "Key Vault e evidências de rotação",
    1544: "Microsoft Defender for Cloud",
    1548: "OWASP Top 10 for LLM Applications",
    1557: "Documentação da OpenAI API",
    1559: "Documentação do Anthropic Claude",
    1628: (
        "Uma organização licencia um serviço empresarial de IA, mas os funcionários continuam usando contas "
        "pessoais porque o fluxo aprovado é mais lento. Assim, prompts sensíveis deixam o ambiente governado."
    ),
    1644: (
        "Um assistente obtém 96 por cento em um conjunto de dados de qualidade, mas revela informações restritas "
        "em dois testes de autorização. A média parece excelente; o sistema não deve ser liberado."
    ),
    1688: (
        "Um inventário dos modelos, dados, código, prompts, ferramentas, serviços e dependências usados por um "
        "sistema de IA."
    ),
    1697: "Chunk",
    1698: "Uma parte de um documento preparada para pesquisa, embedding ou recuperação.",
    1705: "Controle",
    1711: "Embedding",
    1713: "Ajuste fino",
    1715: "Fundamentação",
    1717: "Mecanismo de proteção",
    1729: "LLM",
    1739: "Model card",
    1740: "Uma descrição do uso pretendido, da avaliação, das limitações e de outras informações importantes de um modelo.",
    1743: "Prompt",
    1744: "Texto ou outra entrada fornecida para orientar a resposta de um modelo.",
    1745: "Injeção de prompt",
    1747: "RAG",
    1749: "Red teaming",
    1755: "Security trimming",
    1757: "Service principal",
    1759: "Shadow AI",
    1765: "Tool calling",
    1766: "Solicitação estruturada de uma aplicação de modelo para que um software externo execute uma operação.",
    1773: "Zero Trust",
}


def all_paragraphs(document: Document):
    result = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                result.extend(cell.paragraphs)
    return result


def set_text_preserving_first_run(paragraph, value: str) -> None:
    # Use all descendant text nodes so runs nested in hyperlinks are cleared as
    # well.  python-docx's paragraph.runs omits those hyperlink-contained runs.
    text_nodes = paragraph._p.xpath(".//w:t")
    if text_nodes:
        text_nodes[0].text = value
        for node in text_nodes[1:]:
            node.text = ""
    else:
        paragraph.add_run(value)


def rebuild_table_of_contents(
    english: Document, translated: Document, page_numbers: list[int] | None = None
) -> None:
    english_toc = [p for p in english.paragraphs if p.style.name.casefold() in {"toc 1", "toc 2"}]
    translated_toc = [p for p in translated.paragraphs if p.style.name.casefold() in {"toc 1", "toc 2"}]
    translated_headings = [
        p for p in translated.paragraphs if p.style.name in {"Heading 1", "Heading 2"}
    ]
    if not (len(english_toc) == len(translated_toc) == len(translated_headings) == 207):
        raise SystemExit("TOC/heading alignment failed")

    if page_numbers is not None and len(page_numbers) != len(translated_toc):
        raise SystemExit("TOC page-number map does not contain 207 entries")

    for ordinal, (source, target, heading) in enumerate(
        zip(english_toc, translated_toc, translated_headings)
    ):
        source_links = source._p.xpath("./w:hyperlink")
        if len(source_links) != 1:
            raise SystemExit("authoritative TOC hyperlink structure is unexpected")
        source_text_nodes = source_links[0].xpath(".//w:t")
        if len(source_text_nodes) < 2:
            raise SystemExit("authoritative TOC field result is incomplete")
        page = str(page_numbers[ordinal] if page_numbers else source_text_nodes[-1].text)

        # Google Translate split each TOC entry into two hyperlinks and moved
        # cached page numbers ahead of the heading.  A simple internal hyperlink
        # is deterministic in Word and LibreOffice and keeps navigation intact.
        link = OxmlElement("w:hyperlink")
        for attribute in (qn("w:anchor"), qn("w:history")):
            if source_links[0].get(attribute) is not None:
                link.set(attribute, source_links[0].get(attribute))
        for value, tab in ((heading.text, False), (None, True), (page, False)):
            run = OxmlElement("w:r")
            properties = OxmlElement("w:rPr")
            properties.append(OxmlElement("w:noProof"))
            run.append(properties)
            if tab:
                run.append(OxmlElement("w:tab"))
            else:
                text = OxmlElement("w:t")
                text.text = value
                run.append(text)
            link.append(run)
        ppr = target._p.pPr
        for child in list(target._p):
            if child is not ppr:
                target._p.remove(child)
        target._p.append(link)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--english", type=Path, required=True)
    parser.add_argument("--machine", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--toc-page-map", type=Path)
    args = parser.parse_args()

    english = Document(args.english)
    translated = Document(args.machine)
    english_paragraphs = all_paragraphs(english)
    translated_paragraphs = all_paragraphs(translated)

    if len(english.paragraphs) != len(translated.paragraphs):
        raise SystemExit("body paragraph parity failed")
    if len(english.tables) != len(translated.tables):
        raise SystemExit("table parity failed")
    if len(english_paragraphs) != len(translated_paragraphs):
        raise SystemExit("aggregate paragraph parity failed")
    if [p.style.name for p in english_paragraphs] != [p.style.name for p in translated_paragraphs]:
        raise SystemExit("paragraph style parity failed")
    if [(len(t.rows), len(t.columns)) for t in english.tables] != [
        (len(t.rows), len(t.columns)) for t in translated.tables
    ]:
        raise SystemExit("table shape parity failed")

    for paragraph in translated_paragraphs:
        revised = paragraph.text
        for source, target in GLOBAL_REPLACEMENTS.items():
            revised = revised.replace(source, target)
        if revised != paragraph.text:
            set_text_preserving_first_run(paragraph, revised)

    for index, replacement in OVERRIDES.items():
        if index >= len(translated_paragraphs):
            raise SystemExit(f"correction index {index} is out of range")
        set_text_preserving_first_run(translated_paragraphs[index], replacement)

    page_numbers = None
    if args.toc_page_map:
        page_numbers = [int(value) for value in args.toc_page_map.read_text(encoding="utf-8").split()]
    rebuild_table_of_contents(english, translated, page_numbers)

    translated.core_properties.title = "Proteção de IA Empresarial no Microsoft Azure"
    translated.core_properties.subject = (
        "Guia prático sobre Microsoft Foundry, Purview, Entra, Defender, OpenAI, Claude, "
        "RAG, MCP, agentes, governança e resposta a incidentes"
    )
    translated.core_properties.keywords = (
        "Azure, IA, Microsoft, Foundry, Purview, Entra, Defender, OpenAI, Claude, RAG, "
        "MCP, OWASP, segurança"
    )

    args.output.parent.mkdir(parents=True, exist_ok=True)
    translated.save(args.output)

    # Google Translate rewrote the numbering definitions even though every
    # paragraph retained its authoritative numId. LibreOffice then rendered
    # several list markers over the translated text. Restore numbering.xml
    # byte-for-byte from the English authority so list geometry and glyphs
    # remain source-faithful.
    with zipfile.ZipFile(args.english) as source_package:
        authoritative_numbering = source_package.read("word/numbering.xml")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temporary:
        temporary_path = Path(temporary.name)
    try:
        with zipfile.ZipFile(args.output) as input_package, zipfile.ZipFile(
            temporary_path, "w", zipfile.ZIP_DEFLATED
        ) as output_package:
            for item in input_package.infolist():
                payload = (
                    authoritative_numbering
                    if item.filename == "word/numbering.xml"
                    else input_package.read(item.filename)
                )
                output_package.writestr(item, payload)
        temporary_path.replace(args.output)
    finally:
        temporary_path.unlink(missing_ok=True)


if __name__ == "__main__":
    main()
