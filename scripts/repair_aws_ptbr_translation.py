#!/usr/bin/env python3
"""Repair the source-aligned AWS pt-BR machine-translation first pass.

English remains authoritative. The input must retain exact paragraph and table
alignment; the script exits without publishing if that contract is lost.
"""

from __future__ import annotations

import argparse
import copy
import tempfile
import zipfile
from pathlib import Path

from docx import Document


REPLACEMENTS = {
    "Identity Center do IAM": "IAM Identity Center",
    "Guarda-corpos de rocha matriz da Amazônia": "Amazon Bedrock Guardrails",
    "Centro de Identidade do IAM": "IAM Identity Center",
    "Centro de Identidade IAM": "IAM Identity Center",
}


def all_paragraphs(document: Document):
    result = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                result.extend(cell.paragraphs)
    return result


def set_text_preserving_first_node(paragraph, value: str) -> None:
    nodes = paragraph._p.xpath(".//w:t")
    if nodes:
        nodes[0].text = value
        for node in nodes[1:]:
            node.text = ""
    else:
        paragraph.add_run(value)


def restore_bookmarks(source, target) -> None:
    for node in list(target._p.xpath(".//w:bookmarkStart | .//w:bookmarkEnd")):
        node.getparent().remove(node)
    starts = source._p.xpath("./w:bookmarkStart")
    ends = source._p.xpath("./w:bookmarkEnd")
    insert_at = 1 if target._p.pPr is not None else 0
    for node in starts:
        target._p.insert(insert_at, copy.deepcopy(node))
        insert_at += 1
    for node in ends:
        target._p.append(copy.deepcopy(node))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--english", type=Path, required=True)
    parser.add_argument("--machine", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    english = Document(args.english)
    translated = Document(args.machine)
    source_paragraphs = all_paragraphs(english)
    target_paragraphs = all_paragraphs(translated)

    if len(english.paragraphs) != len(translated.paragraphs):
        raise SystemExit("body paragraph parity failed")
    if len(source_paragraphs) != len(target_paragraphs):
        raise SystemExit("aggregate paragraph parity failed")
    if len(english.tables) != len(translated.tables):
        raise SystemExit("table parity failed")
    if [(len(t.rows), len(t.columns)) for t in english.tables] != [
        (len(t.rows), len(t.columns)) for t in translated.tables
    ]:
        raise SystemExit("table-shape parity failed")
    if len(english.inline_shapes) != len(translated.inline_shapes):
        raise SystemExit("figure parity failed")

    for node in list(translated._element.xpath(".//w:bookmarkStart | .//w:bookmarkEnd")):
        node.getparent().remove(node)

    for source, target in zip(source_paragraphs, target_paragraphs):
        target_ppr = target._p.get_or_add_pPr()
        for existing in list(target_ppr.xpath("./w:pStyle")):
            target_ppr.remove(existing)
        source_style = source._p.xpath("./w:pPr/w:pStyle")
        if source_style:
            target_ppr.insert(0, copy.deepcopy(source_style[0]))
        restore_bookmarks(source, target)
        revised = target.text
        for old, new in REPLACEMENTS.items():
            revised = revised.replace(old, new)
        if revised != target.text:
            set_text_preserving_first_node(target, revised)

    seen_bookmark_names = set()
    for start in list(translated._element.xpath(".//w:bookmarkStart")):
        name = start.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name")
        if name in seen_bookmark_names:
            bookmark_id = start.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
            start.getparent().remove(start)
            for end in list(translated._element.xpath(".//w:bookmarkEnd")):
                if end.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id") == bookmark_id:
                    end.getparent().remove(end)
                    break
        else:
            seen_bookmark_names.add(name)

    translated.core_properties.title = "Proteção da IA Empresarial na Amazon Web Services"
    translated.core_properties.subject = (
        "Guia prático sobre Amazon Bedrock, SageMaker AI, IAM, Macie, Security Hub, "
        "OpenAI, Claude, RAG, MCP, agentes, governança e resposta a incidentes"
    )
    translated.core_properties.keywords = (
        "AWS, IA, Amazon Bedrock, SageMaker AI, IAM, Macie, Security Hub, OpenAI, "
        "Claude, RAG, MCP, OWASP, segurança"
    )
    args.output.parent.mkdir(parents=True, exist_ok=True)
    translated.save(args.output)

    with zipfile.ZipFile(args.english) as source_package:
        numbering = source_package.read("word/numbering.xml")
        styles = source_package.read("word/styles.xml")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temporary:
        temporary_path = Path(temporary.name)
    try:
        with zipfile.ZipFile(args.output) as input_package, zipfile.ZipFile(
            temporary_path, "w", zipfile.ZIP_DEFLATED
        ) as output_package:
            for item in input_package.infolist():
                if item.filename == "word/numbering.xml":
                    payload = numbering
                elif item.filename == "word/styles.xml":
                    payload = styles
                else:
                    payload = input_package.read(item.filename)
                output_package.writestr(item, payload)
        temporary_path.replace(args.output)
    finally:
        temporary_path.unlink(missing_ok=True)


if __name__ == "__main__":
    main()

